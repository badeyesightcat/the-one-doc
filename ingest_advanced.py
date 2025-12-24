import os
# import json
import hashlib
import fitz
import docx
from pathlib import Path
from dateutil import parser as date_parser

def get_file_hash(file_path):
    """Generates a unique MD5 fingerprint for the file content."""
    hasher = hashlib.md5()
    with open(file_path, 'rb') as f:
        buf = f.read()
        hasher.update(buf)
    return hasher.hexdigest()
  
def get_file_metadata(file_path: Path) -> dict:
    """Extracts creation date and author from file properties."""
    stats = os.stat(file_path)
    created_timestamp = getattr(stats, 'st_birthtime', stats.st_ctime)
    meta = {"created_at": created_timestamp, "author": "Unknown"}

    # Try to get specific internal metadata
    if file_path.suffix == ".pdf":
        try:
            with fitz.open(file_path) as doc:
                pdf_meta = doc.metadata
                if pdf_meta.get("creationDate"):
                    date_str = pdf_meta["creationDate"]
                    if date_str.startswith("D:"):
                        date_str = date_str[2:]
                    try:
                        parsed = date_parser.parse(date_str[:14])
                        meta["created_at"] = parsed.timestamp()
                    except ValueError:
                        pass  # Keep filesystem timestamp as fallback
                if pdf_meta.get("author"):
                    meta["author"] = pdf_meta["author"]
        except Exception as e:
            print(f"Warning: Could not extract PDF metadata from {file_path.name}: {e}")

    elif file_path.suffix == ".docx":
        try:
            doc = docx.Document(file_path)
            core_props = doc.core_properties
            if core_props.created:
                meta["created_at"] = core_props.created.timestamp()
            if core_props.author:
                meta["author"] = core_props.author
        except Exception as e:
            print(f"Warning: Could not extract DOCX metadata from {file_path.name}: {e}")

    return meta

def ingest_documents_advanced(folder_path: str) -> list[dict]:
    digital_library = []
    files = list(Path(folder_path).glob("**/*"))
    
    for file_path in files:
        if file_path.suffix not in [".pdf", ".docx"]:
            continue
        
        # 1. Calculate the Hash (The Digital Fingerprint)
        file_hash = get_file_hash(file_path)
        
        # 1. Extract text
        text_content = ""
        try:
            if file_path.suffix == ".pdf":
                with fitz.open(file_path) as doc:
                    for page in doc:
                        text_content += f"{page.get_text()}\n\n"
            elif file_path.suffix == ".docx":
                doc = docx.Document(file_path)
                text_content = "\n\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            print(f"Warning: Failed to extract text from {file_path.name}: {e}")
        
        base_path = Path(folder_path)

        # 3. Create the Document Object
        # We attach 'file_hash' so the next step knows if it can skip AI processing
        doc_entry = {
            "id": str(file_path.relative_to(base_path)),
            "file_hash": file_hash,
            "full_text": text_content,
            "metadata": get_file_metadata(file_path),
            "chunks": []
        }
        
        digital_library.append(doc_entry)
    
    return digital_library

if __name__ == "__main__":
    # Test run
    docs = ingest_documents_advanced("uploads")
    print(f"Ingested {len(docs)} docs. First doc hash: {docs[0].get('file_hash', 'N/A')}")